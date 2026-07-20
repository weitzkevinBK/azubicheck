from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
OUT = DOCS / "docx"

FILES = [
    ("schulleitung-datenschutzpaket.md", "AzubiCheck - Unterlage fuer Schulleitung und Datenschutzpruefung.docx"),
    ("vorlage-verarbeitungsverzeichnis.md", "AzubiCheck - Vorlage Verarbeitungsverzeichnis.docx"),
    ("datenschutzhinweise-entwurf.md", "AzubiCheck - Datenschutzhinweise Entwurf.docx"),
]


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(table, top=80, start=120, bottom=80, end=120):
    tbl_pr = table._tbl.tblPr
    tbl_cell_mar = tbl_pr.first_child_found_in("w:tblCellMar")
    if tbl_cell_mar is None:
        tbl_cell_mar = OxmlElement("w:tblCellMar")
        tbl_pr.append(tbl_cell_mar)
    for margin_name, margin_value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tbl_cell_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tbl_cell_mar.append(node)
        node.set(qn("w:w"), str(margin_value))
        node.set(qn("w:type"), "dxa")


def style_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("AzubiCheck - Stand 20.07.2026")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(90, 90, 90)


def add_title(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string("0B2545")


def add_horizontal_rule(paragraph):
    p = paragraph._p
    p_pr = p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "D9E2EC")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def add_key_value_table(doc, rows):
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    set_cell_margins(table)
    for idx, (key, value) in enumerate(rows):
        left = table.cell(idx, 0)
        right = table.cell(idx, 1)
        left.width = Inches(1.875)
        right.width = Inches(4.625)
        left.text = key
        right.text = value
        set_cell_shading(left, "F2F4F7")
        for paragraph in left.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    doc.add_paragraph()


def add_list_item(doc, text, ordered=False):
    style = "List Number" if ordered else "List Bullet"
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(4)
    p.add_run(text)


def clean_inline(text):
    return re.sub(r"\s+", " ", text).strip()


def build_docx(src, dest):
    lines = src.read_text(encoding="utf-8").splitlines()
    doc = Document()
    style_document(doc)

    pending_table = []
    first_heading = True
    in_codeish_block = False

    def flush_table():
        nonlocal pending_table
        if not pending_table:
            return
        if len(pending_table) >= 2 and pending_table[1].strip().startswith("|---"):
            data_rows = pending_table[2:]
        else:
            data_rows = pending_table
        parsed = []
        for row in data_rows:
            cells = [cell.strip() for cell in row.strip().strip("|").split("|")]
            if len(cells) >= 2:
                parsed.append(cells)
        if parsed:
            table = doc.add_table(rows=len(parsed), cols=max(len(row) for row in parsed))
            table.style = "Table Grid"
            set_cell_margins(table)
            for r_idx, row in enumerate(parsed):
                for c_idx, value in enumerate(row):
                    cell = table.cell(r_idx, c_idx)
                    cell.text = value
                    if r_idx == 0:
                        set_cell_shading(cell, "F2F4F7")
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.bold = True
            doc.add_paragraph()
        pending_table = []

    for raw in lines:
        line = raw.rstrip()
        if line.startswith("|"):
            pending_table.append(line)
            continue
        flush_table()

        if not line:
            continue

        if line.startswith("# "):
            text = clean_inline(line[2:])
            if first_heading:
                add_title(doc, text)
                rule = doc.add_paragraph()
                add_horizontal_rule(rule)
                first_heading = False
            else:
                doc.add_heading(text, level=1)
            continue
        if line.startswith("## "):
            doc.add_heading(clean_inline(line[3:]), level=1)
            continue
        if line.startswith("### "):
            doc.add_heading(clean_inline(line[4:]), level=2)
            continue
        if line.startswith("#### "):
            doc.add_heading(clean_inline(line[5:]), level=3)
            continue
        if line.startswith("- "):
            add_list_item(doc, clean_inline(line[2:]))
            continue
        if re.match(r"^\d+\.\s+", line):
            add_list_item(doc, clean_inline(re.sub(r"^\d+\.\s+", "", line)), ordered=True)
            continue
        if line.endswith(":") and len(line) < 60:
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.bold = True
            continue

        p = doc.add_paragraph()
        p.add_run(clean_inline(line))

    flush_table()
    dest.parent.mkdir(parents=True, exist_ok=True)
    doc.save(dest)


def main():
    OUT.mkdir(exist_ok=True)
    for md_name, docx_name in FILES:
        build_docx(DOCS / md_name, OUT / docx_name)
        print(OUT / docx_name)


if __name__ == "__main__":
    main()
